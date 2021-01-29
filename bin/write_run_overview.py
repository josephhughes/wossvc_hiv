#!/usr/bin/env python3

import argparse
from docx import Document
import datetime
from Bio import SeqIO


# get user arguments
#   ${script_dir}/bin/write_run_overview.py --fasta $BASE --subtypes ${now}_subtypes.txt --tree $TREE_PNG --clusters --reg ${now}.registry.txt

parser = argparse.ArgumentParser()
parser.add_argument('--fasta', required=True, help='input fasta file with new samples')
parser.add_argument('--subtypes', required=True, help='text file with subtype information for all samples')
parser.add_argument('--tree', required=True, help='pdf of phylogenetic tree generated from input samples')
parser.add_argument('--clusters', required=False, help='report will include details of clustering information updated after the current run')
parser.add_argument('--reg',required=False,help='new registry which will be used for determining the clusters of the new sequences')
args = parser.parse_args()


fastaIn = ''
subtypes=''
tree=''

# set user arguments
if args.fasta:
    fastaIn = args.fasta
if args.subtypes:
    subtypes = args.subtypes
if args.tree:
    tree = args.tree

now = datetime.datetime.now()
time_now = (str('{:02d}'.format(now.day)) + "-" +str('{:02d}'.format(now.month)) + "-" + str(now.year))


# overview of new samples
new_samples = {}

for record in SeqIO.parse(fastaIn, 'fasta'):
    new_samples[record.id] = record.seq


# subtype overview
subtype_counts = {}

with open(subtypes, 'r') as f:
    for line in f:
        values = line.split("\t")

        sample = values[0]
        subtype = values[1].split(" ")[0]

        if subtype in subtype_counts.keys():
            subtype_counts[subtype] +=1
        else:
            subtype_counts[subtype] = 1


# Add cluster run details?
# e.g. number of samples currently in the database
# detailed cluster information
# changes & important updates?



# CREATE REPORT
outFile = time_now + "_RunOverview.docx"
document = Document()
document.add_heading("HIV Sample Analysis Run Summary\n", 0)

p1 = document.add_paragraph("Date of run:\t" + time_now + "\nNumber of samples added:\t" + str(len(new_samples)))

p2 = document.add_paragraph("Number of each subtype:\n")
for key, value in subtype_counts.items():
    p2.add_run("\t\t" + key).bold=True
    p2.add_run(":\t" + str(value) + "\n")

document.add_picture(tree)


# if cluster information is included - add that here
# parse registry and provide the cluster number that they are in
if args.clusters and args.reg:
    document.add_heading("Cluster Analysis\n", level=1)
    p3 = document.add_paragraph("Cluster assignment for subtype C samples processed with HIVtrace:\n")
    with open(args.reg, 'r') as rf:
      for line in rf:
        values = line.split("\t")
        if values[0] in new_samples:
          #print(values[0] + values[-1])
          p3.add_run("\t" + values[0]).bold=True
          p3.add_run("\t" + values[-1])
    with open(args.clusters,'r') as cf:
      p4 = document.add_paragraph("Cluster changes:\n")
      for line in cf:
        p4.add_run(line)
       

document.save(outFile)
